from openai import OpenAI
client = OpenAI()
import docx
from io import BytesIO
from PIL import Image
import os
IMAGE_DIR = './tempImages/'
SERVER_URL = 'https://assistant.arpasistemas.com.br/api/getTempImage'



def get_image_description(image_stream, image_name):
    try:
        image = Image.open(image_stream)
        temp_image_path = os.path.join(IMAGE_DIR, image_name)
        image.save(temp_image_path)
        image_url = f"{SERVER_URL}/{image_name}"
        user_prompt = "pelase, make a short description of this image in portuguese-br. it will be used on a documentation to be used by an gpt assistant"
        system_prompt  = (
            "Please try to generate the text for an user manual documentation based on and assistat."
        )
        response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
            "role": "user",
            "content": [
                {"type": "text", "text": user_prompt},
                {
                "type": "image_url",
                "image_url": {
                    "url": image_url
                },
                },
            ],
            }
        ],
        max_tokens=300,
        )
        return response.choices[0].message.content
    except Exception as error:
        print("Error processing image:", error)
        raise error
def replace_images_with_text(doc_path):
    doc = docx.Document(doc_path)
    new_doc = docx.Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                inline_shapes = run._element.xpath('.//w:drawing')
                for shape in inline_shapes:
                    image_data = shape.xpath('.//a:blip/@r:embed')
                    if image_data:
                        image_part = doc.part.related_parts[image_data[0]]
                        image_stream = BytesIO(image_part.blob)
                        image_name = f"temp_image_{image_data[0]}.png"  # Generate a unique name
                        description = get_image_description(image_stream, image_name)
                        new_paragraph.add_run(description)
            else:
                new_paragraph.add_run(run.text)
    new_doc.save("output.docx")
    print("Imangns substituidas pela descricao. novo arquivo: output.docx")
replace_images_with_text("/Users/programacao/dev/gpt/src/docx/smartv5.docx")
