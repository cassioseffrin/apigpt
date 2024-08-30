import sys
import os
import sqlite3
from docx import Document
from xml.etree import ElementTree as ET
from openai import OpenAI
SERVER_URL = 'https://assistant.arpasistemas.com.br/api/getTempImage'
client = OpenAI()
def setup_database(db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS assistant (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            assistantId TEXT UNIQUE NOT NULL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS images (
            filename TEXT NOT NULL,
            title TEXT,
            description TEXT,
            assistant_id INTEGER,
            FOREIGN KEY (assistant_id) REFERENCES assistant(id)
        )
    ''')
    conn.commit()
    return conn

 
def get_description_from_db(conn, img_caption):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT title, description FROM images WHERE filename = ?
    ''', (img_caption,))
    result = cursor.fetchone()
    if result:
        title, description = result
        return title, description
    else:
        return "sem titulo", "sem descricao"


def get_filename_from_title_db(conn, title):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT filename FROM images WHERE title = ? LIMIT 1
    ''', (title,))
    result = cursor.fetchone()
    if result:
        return result[0]
    else:
        return ""
def extract_images_from_docx(document):
    image_data = []
    for shape in document.inline_shapes:
        graphic = shape._inline.graphic
        if not graphic:
            continue
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        namespaces = {
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            alt_text = cNvPr_element.get('descr', None)
            if not alt_text:
                alt_text = cNvPr_element.get('name', None)
            if not alt_text:
                alt_text = f'image_{len(image_data)+1:04d}'
            blip_element = pic_element.find('.//a:blip', namespaces)
            if blip_element is not None:
                embed_id = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                rel = document.part.rels[embed_id]
                if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                    image = rel.target_part.blob
                    filename = f'image_{embed_id}.png'
                    image_data.append((image, alt_text, filename, shape))
    return image_data
def save_images_to_disk(image_data, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    for i, (img, img_desc, img_caption, _) in enumerate(image_data):
        img_path = os.path.join(output_dir, f'{img_caption}')
        with open(img_path, 'wb') as img_file:
            img_file.write(img)
        print(f'Saved image to {img_path}')
def add_image_description_to_docx(doc_path, output_path, conn):
    doc = Document(doc_path)
    image_data = extract_images_from_docx(doc)
    for image, alt_text, filename, shape in image_data:
        description = get_description_from_db(conn, filename)
        graphic = shape._inline.graphic
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        namespaces = {
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            cNvPr_element.set('descr', description)
            cNvPr_element.set('title', description)
        parent_paragraph = shape._inline.getparent().getparent().getparent()
        parent_element = parent_paragraph.getparent()
        index = parent_element.index(parent_paragraph)
        new_paragraph = doc.add_paragraph(f"{description} IMAGE_FILENAME: ({filename})")
        parent_element.insert(index + 1, new_paragraph._element)
    doc.save(output_path)
    print(f"Images updated and descriptions added. New file: {output_path}")
def get_image_description(image_name):
    try:
        image_url = f"{SERVER_URL}/{image_name}"
        user_prompt = "Por favor, faça uma breve descrição desta imagem em português-br. Será usada em uma documentação para ser usada por um assistente GPT."
        system_prompt = "Please try to generate the text for an user manual documentation based on an assistant."
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {"type": "image_url", "image_url": {"url": image_url}},
                    ],
                }
            ],
            max_tokens=300,
        )
        return response.choices[0].message.content
    except Exception as error:
        print("Error processing image:", error)
        raise error
def insert_image_data(conn, image_data, assistant_name, assistant_id, updateAiDescription):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR IGNORE INTO assistant (name, assistantId) VALUES (?, ?)
    ''', (assistant_name, assistant_id))
    cursor.execute('''
        SELECT id FROM assistant WHERE assistantId = ?
    ''', (assistant_id,))
    assistant_row_id = cursor.fetchone()[0]
    for img, img_title, img_caption, _ in image_data:  
        img_description = None
        if updateAiDescription:
            img_description = get_image_description(f'{img_caption}')
            # img_description += "filename: f{filename}"
        cursor.execute('''
            INSERT INTO images (filename, title, description, assistant_id) VALUES (?, ?, ?, ?)
        ''', (f'{img_caption}', img_title, img_description, assistant_row_id))
        conn.commit()  
def cleanup_files(conn, output_dir, assistant_id):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT images.filename FROM images
        JOIN assistant ON images.assistant_id = assistant.id
        WHERE assistant.assistantId = ?
    ''', (assistant_id,))
    files_to_delete = cursor.fetchall()
    for (filename,) in files_to_delete:
        img_path = os.path.join(output_dir, filename)
        if os.path.exists(img_path):
            os.remove(img_path)
            print(f'Deleted image {img_path}')
    cursor.execute('''
        DELETE FROM images
        WHERE assistant_id = (SELECT id FROM assistant WHERE assistantId = ?)
    ''', (assistant_id,))
    conn.commit()

def replace_images_with_text(doc_path, output_path, conn):
    doc = Document(doc_path)
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                inline_shapes = run._element.xpath('.//w:drawing')
                for shape in inline_shapes:
                    image_data = shape.xpath('.//a:blip/@r:embed')
                    if image_data:
                        embed_id = image_data[0]
                        img_caption = f'image_{embed_id}.png'
                        # description = get_description_from_db(conn, img_caption)
                        title, description = get_description_from_db(conn, img_caption)
                        new_paragraph.add_run(f"{title} IMAGE_FILENAME: ({img_caption}), Descrição: ")
                        new_paragraph.add_run(f"{description} IMAGE_FILENAME: ({img_caption})")
            else:
                if (run.text.startswith("Figura") or run.text.startswith("Fonte: Aplicativo Play")):
                    run.text=""
                    # filename = get_filename_from_title_db(conn, run.text)
                    # run.text = f"{paragraph.text} IMAGE_FILENAME: ({filename})"
                new_paragraph.add_run(run.text)
    new_doc.save(output_path)
    print(f"Images replaced with descriptions. New file: {output_path}")
 



def replace_images_with_text2(doc_path, output_path, conn):
    doc = Document(doc_path)
    # new_doc = Document()

    figura_paragraphs = {}  # Dictionary to store "Figura*" paragraphs and their text

    # First pass: Store paragraphs starting with "Figura*"
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("Figura"):
            paragraph.text = f"{paragraph.text} teste"

    # Second pass: Process paragraphs and add text inline
    # for i, paragraph in enumerate(doc.paragraphs):
    #     new_paragraph = new_doc.add_paragraph()
    #     if i in figura_paragraphs:
    #         # Add the "Figura*" text
    #         new_paragraph.add_run(f"{figura_paragraphs[i]}")
        
    #     for run in paragraph.runs:
    #         if run._element.xpath('.//w:drawing'):
    #             inline_shapes = run._element.xpath('.//w:drawing')
    #             for shape in inline_shapes:
    #                 image_data = shape.xpath('.//a:blip/@r:embed')
    #                 if image_data:
    #                     embed_id = image_data[0]
    #                     img_caption = f'image_{embed_id}.png'
    #                     # Append the image filename to the same paragraph
    #                     new_paragraph.add_run(f" image_filename: ({img_caption})")
    #         else:
    #             new_paragraph.add_run(run.text)
    
    doc.save(output_path)
    print(f"Images replaced with descriptions. New file: {output_path}")





def main():
    if len(sys.argv) < 5:
        print("Usage: python extract_images.py <file_path> <output_dir> <assistant_id> <cleanup> <updateAiDescription>")
        sys.exit(1)
    file_path = sys.argv[1]
    output_dir = sys.argv[2]
    assistant_id = sys.argv[3]
    cleanup = sys.argv[4].lower() == 'true'
    updateAiDescription = sys.argv[5].lower() == 'true'
    db_path = 'images_assistant.db'
    conn = setup_database(db_path)
    if cleanup:
        cleanup_files(conn, output_dir, assistant_id)
    doc = Document(file_path)
    image_data = extract_images_from_docx(doc)
    save_images_to_disk(image_data, output_dir)
    insert_image_data(conn, image_data, 'Smart Vendas', assistant_id, updateAiDescription)
    output_path = f"{os.path.splitext(file_path)[0]}_com_descricao.docx"
    output_path_without_images = f"{os.path.splitext(file_path)[0]}_vector_store.docx"
    if updateAiDescription:
        add_image_description_to_docx(file_path, output_path, conn)
    replace_images_with_text(file_path, output_path_without_images, conn)
    conn.close()
if __name__ == "__main__":
    main()
