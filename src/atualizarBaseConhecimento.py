import sys
import os
import sqlite3
from docx import Document
from xml.etree import ElementTree as ET
from openai import OpenAI
import re
import unicodedata
SERVER_URL = 'https://assistant.arpasistemas.com.br/api/getImage'
# SERVER_URL = 'http://127.0.0.1:4014/api/getImage'
client = OpenAI()
def setup_database(db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS assistant (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            assistantId TEXT UNIQUE NOT NULL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS images (
            filepath TEXT NOT NULL,
            filename TEXT NOT NULL,
            content_type TEXT NOT NULL,
            title TEXT,
            description TEXT,
            assistant_id INTEGER,
            FOREIGN KEY (assistant_id) REFERENCES assistant(id)
        )
    ''')
    conn.commit()
    return conn
def get_description_from_db(conn, img_caption):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT title, description FROM images WHERE filename = ?
    ''', (img_caption,))
    result = cursor.fetchone()
    if result:
        title, description = result
        return title, description
    else:
        return "sem titulo", "sem descricao"
def get_filename_from_title_db(conn, title):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT filename FROM images WHERE title = ? LIMIT 1
    ''', (title,))
    result = cursor.fetchone()
    if result:
        return result[0]
    else:
        return ""
def extract_images_from_docx(document):
    image_data = []
    for shape in document.inline_shapes:
        graphic = shape._inline.graphic
        if not graphic:
            continue
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        namespaces = {
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            alt_text = cNvPr_element.get('descr', None)
            if not alt_text:
                alt_text = cNvPr_element.get('name', None)
            if not alt_text:
                alt_text = f'image_{len(image_data)+1:04d}'
            blip_element = pic_element.find('.//a:blip', namespaces)
            if blip_element is not None:
                embed_id = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                rel = document.part.rels[embed_id]
                if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                    image = rel.target_part.blob
                    content_type = rel._target.image.content_type
                    ext = rel._target.image.ext
                    if ext == None or len(ext) == 0:
                        ext = "png"
                    filename = generate_filename_from_alt_text(alt_text,ext)
                    image_data.append((image, alt_text, filename, shape, content_type))
    return image_data
def generate_filename_from_alt_text(alt_text,   ext):
    alt_text = alt_text.strip().lower()
    figura_match = re.match(r'figura\s*(\d+)', alt_text)
    if figura_match:
        figure_number = int(figura_match.group(1))
        return f"figura{figure_number:02d}.{ext}"
    picture_match = re.match(r'picture\s*(\d+)', alt_text)
    if picture_match:
        picture_number = int(picture_match.group(1))
        return f"picture{picture_number}.{ext}"
    alt_text_no_spaces = ''.join(alt_text.split())
    alt_text_normalized = ''.join(
        c for c in unicodedata.normalize('NFD', alt_text_no_spaces)
        if unicodedata.category(c) != 'Mn'
    )
    # return f"{alt_text_normalized}.{ext}"
    #arquivo gerado pelo google docs ja traz extensao
    return f"{alt_text_normalized}"
def save_images_to_disk(image_data, vectorStorePath):
    completeFilePath = "./src/imgs/"+vectorStorePath
    if not os.path.exists(completeFilePath):
        os.makedirs(completeFilePath)
    for i, (img, img_desc, img_caption, content_type, _) in enumerate(image_data):
        img_path = os.path.join(completeFilePath, f'{vectorStorePath}_{img_caption}')
        with open(img_path, 'wb') as img_file:
            img_file.write(img)
        print(f"Imagem salva em {img_path}")
def add_image_description_to_docx(doc_path, output_path, conn):
    doc = Document(doc_path)
    image_data = extract_images_from_docx(doc)
    for image, alt_text, filename, shape, _ in image_data:
        description = get_description_from_db(conn, filename)
        graphic = shape._inline.graphic
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        namespaces = {
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            cNvPr_element.set('descr', description)
            cNvPr_element.set('title', description)
        parent_paragraph = shape._inline.getparent().getparent().getparent()
        parent_element = parent_paragraph.getparent()
        index = parent_element.index(parent_paragraph)
        new_paragraph = doc.add_paragraph(f"{description} IMAGE_FILENAME: ({filename})")
        parent_element.insert(index + 1, new_paragraph._element)
    doc.save(output_path)
    print(f"Imagens atualizadas e descrições adicionadas. Novo arquivo: {output_path}")
def get_image_description(image_path, image_name):
    try:
        image_url = f"{SERVER_URL}/{image_path}_{image_name}"
        user_prompt = "Por favor, faça uma breve descrição desta imagem em português-br. Será usada em uma documentação para ser usada por um assistente GPT."
        system_prompt = "Please try to generate the text for an user manual documentation based on an assistant."
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {"type": "image_url", "image_url": {"url": image_url}},
                    ],
                }
            ],
            max_tokens=300,
        )
        # res = response.choices[0].message.content
        # print(res)
        return response.choices[0].message.content
    except Exception as error:
        print("Erro ao processar imagem:", error)
        raise error
def insert_image_data(conn, image_data, assistant_name, assistant_id, updateAiDescription, filepath):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR IGNORE INTO assistant (name, assistantId) VALUES (?, ?)
    ''', (assistant_name, assistant_id))
    cursor.execute('''
        SELECT id FROM assistant WHERE assistantId = ?
    ''', (assistant_id,))
    assistant_row_id = cursor.fetchone()[0]
    for img, img_title, img_caption, _, content_type in image_data:  
        img_description = None
        if updateAiDescription:
            img_description = get_image_description(filepath, f'{img_caption}')
            # img_description += "filename: f{filename}"
        cursor.execute('''
            INSERT INTO images (filepath, filename, content_type, title, description, assistant_id) VALUES (?, ?, ?, ?, ?, ?)
        ''', (filepath, f'{filepath}_{img_caption}', content_type, img_title, img_description, assistant_row_id))
        conn.commit()  
def cleanup_files(conn, filepath, assistant_id):
    cursor = conn.cursor()
    cursor.execute('''
        DELETE FROM images
        WHERE assistant_id = (SELECT id FROM assistant WHERE assistantId = ? AND filepath = ? )
    ''', (assistant_id, filepath))
    conn.commit()
def replace_images_with_text_ignore_strings(doc_path, filepath, output_path, conn):
    doc = Document(doc_path)
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                inline_shapes = run._element.xpath('.//w:drawing')
                for shape in inline_shapes:
                    image_data = shape.xpath('.//a:blip/@r:embed')
                    if image_data:
                        embed_id = image_data[0]
                        img_caption = f'image_{embed_id}.png'
                        # description = get_description_from_db(conn, img_caption)
                        title, description = get_description_from_db(conn, img_caption)
                        new_paragraph.add_run(f"{title} IMAGE_FILENAME: ({filepath}_{img_caption}), Descrição: ")
                        new_paragraph.add_run(f"{description} IMAGE_FILENAME: ({filepath}_{img_caption})")
            else:
                if (run.text.startswith("Figura") or run.text.startswith("Fonte: Aplicativo")):
                    run.text=""
                    # filename = get_filename_from_title_db(conn, run.text)
                    # run.text = f"{paragraph.text} IMAGE_FILENAME: ({filename})"
                else:
                    new_paragraph.add_run(run.text)
    new_doc.save(output_path)
    print(f"Imagens substituídas por descrições. Novo arquivo: {output_path}")
def replace_images_with_text_old(doc_path, filepath, output_path, conn):
    doc = Document(doc_path)
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                inline_shapes = run._element.xpath('.//w:drawing')
                for shape in inline_shapes:
                    image_data = shape.xpath('.//a:blip/@r:embed')
                    if image_data:
                        embed_id = image_data[0]
                        img_caption = f'image_{embed_id}.png'
                        complete_filename=f"{filepath}_{img_caption}"
                        title, description = get_description_from_db(conn, complete_filename)
                        new_paragraph.add_run(f"{title} IMAGE_FILENAME: ({complete_filename}), Descrição: ")
                        new_paragraph.add_run(f"{description} IMAGE_FILENAME: ({complete_filename})")
            else:
                new_paragraph.add_run(run.text)
    new_doc.save(output_path)
    print(f"Imagens substituídas por descrições. Novo arquivo: {output_path}")
def replace_images_with_text(image_data, doc_path, filepath, output_path, conn):
    doc = Document(doc_path)
    new_doc = Document()
    image_index = 0  
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()   
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):   
                if image_index < len(image_data):
                    image, alt_text, filename, shape, content_type = image_data[image_index]
                    image_filename = f"{filepath}_{filename}"
                    title, description = get_description_from_db(conn, image_filename)
                    # new_paragraph.add_run(f"{title} IMAGE_FILENAME: ({image_filename}), Descrição: ")
                    new_paragraph.add_run(f"{description} IMAGE_URI: \"{image_filename}\".")
                    image_index += 1
            else:
                new_paragraph.add_run(run.text)
    # Save the new document after replacing images with text
    new_doc.save(output_path)
    print(f"Imagens substituídas por descrições. Novo arquivo: {output_path}")
def main():
    if len(sys.argv) < 5:
        print("Uso: python atualizarBaseConhecimento.py <nome_do_arquivo_de_armazenamento_de_vetores> <caminho_do_arquivo> <id_do_assistente> <limpar_base_dados> <atualizar_descricao_ai>")
        sys.exit(1)
    vector_store_filename = sys.argv[1]
    filepath = sys.argv[2]
    assistant_id = sys.argv[3]
    cleanup = sys.argv[4].lower() == 'true'
    updateAiDescription = sys.argv[5].lower() == 'true'
    db_path = 'images_assistant.db'
    conn = setup_database(db_path)
    doc = Document(vector_store_filename)
    image_data = extract_images_from_docx(doc)
    save_images_to_disk(image_data, filepath)
    if cleanup:
        cleanup_files(conn, filepath, assistant_id)
        if (updateAiDescription==False):
            insert_image_data(conn, image_data, 'Smart Vendas', assistant_id, updateAiDescription, filepath)
    if updateAiDescription:
        insert_image_data(conn, image_data, 'Smart Vendas', assistant_id, updateAiDescription, filepath)
    output_path_without_images = f"{os.path.splitext(vector_store_filename)[0]}_data.docx"
    if updateAiDescription:
        add_image_description_to_docx(vector_store_filename, filepath, conn)
    replace_images_with_text(image_data, vector_store_filename, filepath, output_path_without_images, conn)
    conn.close()
if __name__ == "__main__":
    main()