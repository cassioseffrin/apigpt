import sys
import os
import unicodedata
import re
import sqlite3
from docx import Document
from PIL import Image
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import base64
from xml.etree import ElementTree as ET
from openai import OpenAI
from io import BytesIO

IMAGE_DIR = './tempImages/'
SERVER_URL = 'https://assistant.arpasistemas.com.br/api/getTempImage'
client = OpenAI()

def sanitize_filename(text):
    text = unicodedata.normalize('NFKD', text)
    text = text.encode('ascii', 'ignore').decode('ascii')
    text = re.sub(r'[^\w\s-]', '', text).strip().lower()
    text = re.sub(r'[-\s]+', '_', text)
    return text

def get_image_description(image_name):
    try:
        image_url = f"{SERVER_URL}/{image_name}"
        user_prompt = "Por favor, faça uma breve descrição desta imagem em português-br. Será usada em uma documentação para ser usada por um assistente GPT."
        system_prompt = "Please try to generate the text for an user manual documentation based on an assistant."
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {"type": "image_url", "image_url": {"url": image_url}},
                    ],
                }
            ],
            max_tokens=300,
        )
        return response.choices[0].message.content
    except Exception as error:
        print("Error processing image:", error)
        raise error

def extract_images_from_docx(file_path):
    document = Document(file_path)
    image_data = []
    for shape in document.inline_shapes:
        graphic = shape._inline.graphic
        if not graphic:
            continue
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        namespaces = {
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            alt_text = cNvPr_element.get('descr', None)
            if not alt_text:
                alt_text = cNvPr_element.get('name', None)
            if not alt_text:
                alt_text = f'image_{len(image_data)+1:04d}'
            blip_element = pic_element.find('.//a:blip', namespaces)
            if blip_element is not None:
                embed_id = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                rel = document.part.rels[embed_id]
                if rel.reltype == RT.IMAGE:
                    image = rel.target_part.blob
                    filename = f'image_{embed_id}.png'
                    image_data.append((image, alt_text, filename))
    return image_data, document

import xml.etree.ElementTree as ET

def print_xml(element, message=""):
    print(f"{message}:\n{ET.tostring(element, encoding='unicode')}\n")

def update_images_alt_text_with_description(doc_path, output_path, conn):
    doc = Document(doc_path)
    namespaces = {
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }

    for shape in doc.inline_shapes:
        graphic = shape._inline.graphic
        if not graphic:
            continue
        graphic_xml = graphic.xml
        pic_element = ET.fromstring(graphic_xml)
        
        cNvPr_element = pic_element.find('.//pic:cNvPr', namespaces)
        if cNvPr_element is not None:
            print_xml(cNvPr_element, "Before Update")
            
            existing_alt_text = cNvPr_element.get('descr', None) or cNvPr_element.get('name', None) or f'image_{shape._inline.graphic.graphicData.uri}'
            
            blip_element = pic_element.find('.//a:blip', namespaces)
            if blip_element is not None:
                embed_id = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                img_caption = f'image_{embed_id}.png'
                description = get_description_from_db(conn, img_caption)
                if description:
                    updated_alt_text = f"{existing_alt_text}\n\n{description}"
                    
                    # Directly manipulate the underlying XML tree
                    cNvPr_element.attrib['descr'] = updated_alt_text
                    cNvPr_element.attrib['name'] = updated_alt_text
                    cNvPr_element.attrib['title'] = updated_alt_text

                    # Print the XML after the update
                    print_xml(cNvPr_element, "After Update")

                    # Update the XML element in the shape
                    shape._inline.graphic._element = pic_element

    doc.save(output_path)
    print(f"Updated alt text with descriptions. New file: {output_path}")




def save_images_to_disk(image_data, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    for i, (img, img_desc, img_caption) in enumerate(image_data):
        img_path = os.path.join(output_dir, f'{img_caption}.png')
        with open(img_path, 'wb') as img_file:
            img_file.write(img)
        print(f'Saved image to {img_path}')

def encode_images_to_base64(image_data):
    return [(img_caption, base64.b64encode(img).decode('utf-8')) for img, img_desc, img_caption in image_data]

def setup_database(db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS assistant (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            assistantId TEXT UNIQUE NOT NULL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS images (
            filename TEXT NOT NULL,
            title TEXT,
            description TEXT,
            assistant_id INTEGER,
            FOREIGN KEY (assistant_id) REFERENCES assistant(id)
        )
    ''')
    conn.commit()
    return conn

def insert_image_data(conn, image_data, assistant_name, assistant_id, updateAiDescription):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR IGNORE INTO assistant (name, assistantId) VALUES (?, ?)
    ''', (assistant_name, assistant_id))
    cursor.execute('''
        SELECT id FROM assistant WHERE assistantId = ?
    ''', (assistant_id,))
    assistant_row_id = cursor.fetchone()[0]
    for img, img_title, img_caption in image_data:
        img_description = None
        if updateAiDescription:
            img_description = get_image_description(f'{img_caption}.png')
        cursor.execute('''
            INSERT INTO images (filename, title, description, assistant_id) VALUES (?, ?, ?, ?)
        ''', (f'{img_caption}', img_title, img_description, assistant_row_id))
        conn.commit() # comitando a cada registro, evitar perder tudo se der erro no meio
    #conn.commit()

def cleanup_files(conn, output_dir, assistant_id):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT images.filename FROM images
        JOIN assistant ON images.assistant_id = assistant.id
        WHERE assistant.assistantId = ?
    ''', (assistant_id,))
    files_to_delete = cursor.fetchall()
    for (filename,) in files_to_delete:
        img_path = os.path.join(output_dir, filename)
        if os.path.exists(img_path):
            os.remove(img_path)
            print(f'Deleted image {img_path}')
    cursor.execute('''
        DELETE FROM images
        WHERE assistant_id = (SELECT id FROM assistant WHERE assistantId = ?)
    ''', (assistant_id,))
    conn.commit()

 

def get_description_from_db(conn, img_caption):
    cursor = conn.cursor()
    cursor.execute('''
        SELECT description FROM images WHERE filename = ?
    ''', (img_caption,))
    result = cursor.fetchone()
    if result:
        return result[0]
    else:
        return "sem descricao"

def replace_images_with_text(doc_path, output_path, conn):
    doc = Document(doc_path)
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                inline_shapes = run._element.xpath('.//w:drawing')
                for shape in inline_shapes:
                    image_data = shape.xpath('.//a:blip/@r:embed')
                    if image_data:
                        embed_id = image_data[0]
                        img_caption = f'image_{embed_id}.png'
                        description = get_description_from_db(conn, img_caption)
                        new_paragraph.add_run(description)
            else:
                new_paragraph.add_run(run.text)
    new_doc.save(output_path)
    print(f"Images replaced with descriptions. New file: {output_path}")


def main():
    if len(sys.argv) < 5:
        print("Usage: python extract_images.py <file_path> <output_dir> <assistant_id> <cleanup> <updateAiDescription>")
        sys.exit(1)

    file_path = sys.argv[1]
    output_dir = sys.argv[2]
    assistant_id = sys.argv[3]
    cleanup = sys.argv[4].lower() == 'true'
    updateAiDescription = sys.argv[5].lower() == 'true'
    db_path = 'images_assistant.db'
    conn = setup_database(db_path)
    if cleanup:
        cleanup_files(conn, output_dir, assistant_id)

    # image_data, document = extract_images_from_docx(file_path)
    # save_images_to_disk(image_data, output_dir)
    # insert_image_data(conn, image_data, 'Smart Vendas', assistant_id, updateAiDescription)

    if updateAiDescription:
        # replace_images_with_text(file_path, f"{os.path.splitext(file_path)[0]}_without_images.docx", conn)
        update_images_alt_text_with_description(file_path, f"{os.path.splitext(file_path)[0]}_updated_description.docx", conn)
 

    conn.close()

if __name__ == "__main__":
    main()
